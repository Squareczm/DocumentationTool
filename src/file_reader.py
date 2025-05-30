"""
文件内容读取器
支持读取不同格式文件的内容：.docx, .xlsx, .pdf, .txt, .md
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

# 导入文件处理库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class FileReader:
    """文件内容读取器"""
    
    def __init__(self):
        """初始化文件读取器"""
        self.logger = logging.getLogger(__name__)
    
    def read_file(self, file_path: Path) -> Dict[str, Any]:
        """
        读取文件内容和元数据
        
        Args:
            file_path: 文件路径
            
        Returns:
            包含文件内容和元数据的字典
        """
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 获取文件基本信息
        stat = file_path.stat()
        file_info = {
            'path': str(file_path),
            'name': file_path.name,
            'stem': file_path.stem,
            'suffix': file_path.suffix.lower(),
            'size': stat.st_size,
            'creation_time': datetime.fromtimestamp(stat.st_ctime),
            'modification_time': datetime.fromtimestamp(stat.st_mtime),
            'content': '',
            'metadata': {}
        }
        
        # 根据文件扩展名选择合适的读取方法
        try:
            if file_info['suffix'] == '.docx':
                file_info.update(self._read_docx(file_path))
            elif file_info['suffix'] == '.xlsx':
                file_info.update(self._read_xlsx(file_path))
            elif file_info['suffix'] == '.pdf':
                file_info.update(self._read_pdf(file_path))
            elif file_info['suffix'] in ['.txt', '.md']:
                file_info.update(self._read_text(file_path))
            else:
                raise ValueError(f"不支持的文件类型: {file_info['suffix']}")
                
        except Exception as e:
            self.logger.error(f"读取文件内容失败 {file_path}: {e}")
            raise
        
        return file_info
    
    def _read_docx(self, file_path: Path) -> Dict[str, Any]:
        """读取Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("未安装 python-docx 库，无法读取 .docx 文件")
        
        try:
            doc = Document(file_path)
            
            # 提取文本内容
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(' | '.join(row_text))
            
            # 提取文档属性
            core_props = doc.core_properties
            metadata = {}
            
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['created'] = core_props.created
            if core_props.modified:
                metadata['modified'] = core_props.modified
            
            return {
                'content': '\n'.join(content),
                'metadata': metadata
            }
            
        except Exception as e:
            self.logger.error(f"读取Word文档失败: {e}")
            raise
    
    def _read_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """读取Excel文档"""
        if not XLSX_AVAILABLE:
            raise ImportError("未安装 openpyxl 库，无法读取 .xlsx 文件")
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            content = []
            metadata = {}
            
            # 提取工作表内容
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"工作表: {sheet_name}")
                
                # 读取有数据的行
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else '' for cell in row]
                    if any(cell.strip() for cell in row_data):
                        content.append(' | '.join(row_data))
            
            # 提取文档属性
            props = workbook.properties
            if props.title:
                metadata['title'] = props.title
            if props.creator:
                metadata['author'] = props.creator
            if props.subject:
                metadata['subject'] = props.subject
            if props.created:
                metadata['created'] = props.created
            if props.modified:
                metadata['modified'] = props.modified
            
            workbook.close()
            
            return {
                'content': '\n'.join(content),
                'metadata': metadata
            }
            
        except Exception as e:
            self.logger.error(f"读取Excel文档失败: {e}")
            raise
    
    def _read_pdf(self, file_path: Path) -> Dict[str, Any]:
        """读取PDF文档"""
        if not PDF_AVAILABLE:
            raise ImportError("未安装 PyPDF2 或 pdfplumber 库，无法读取 .pdf 文件")
        
        content = []
        metadata = {}
        
        # 首先尝试使用pdfplumber（更好的文本提取）
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
                
                # 提取元数据
                if pdf.metadata:
                    metadata.update({
                        k.replace('/', ''): v for k, v in pdf.metadata.items() if v
                    })
        except Exception as e:
            self.logger.warning(f"pdfplumber提取失败，尝试PyPDF2: {e}")
            
            # 回退到PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    
                    # 提取文本
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            content.append(text)
                    
                    # 提取元数据
                    if reader.metadata:
                        metadata.update({
                            k.replace('/', ''): v for k, v in reader.metadata.items() if v
                        })
                        
            except Exception as e:
                self.logger.error(f"PyPDF2提取也失败: {e}")
                raise
        
        return {
            'content': '\n'.join(content),
            'metadata': metadata
        }
    
    def _read_text(self, file_path: Path) -> Dict[str, Any]:
        """读取纯文本文件"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise UnicodeDecodeError("无法使用常见编码读取文件")
            
            return {
                'content': content,
                'metadata': {'encoding': encoding}
            }
            
        except Exception as e:
            self.logger.error(f"读取文本文件失败: {e}")
            raise
    
    @staticmethod
    def get_supported_extensions() -> list[str]:
        """获取支持的文件扩展名"""
        extensions = ['.txt', '.md']
        
        if DOCX_AVAILABLE:
            extensions.append('.docx')
        if XLSX_AVAILABLE:
            extensions.append('.xlsx')
        if PDF_AVAILABLE:
            extensions.append('.pdf')
        
        return extensions 